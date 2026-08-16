"""
FraudGuard SA - Fraudulent Job Posting Detection System
Flask backend with Random Forest classifier (TF-IDF features).

Accepts a job posting two ways:
  1. Pasted text (JSON body: {"text": "..."})
  2. Uploaded file (multipart form field "file": .txt, .pdf or .docx)
Both paths feed the same validation and classification pipeline.
"""
import io
import os
import pickle
from pathlib import Path
from flask import Flask, request, jsonify, send_file

BASE_DIR = Path(__file__).resolve().parent

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 5 * 1024 * 1024  # 5 MB upload cap

with open(BASE_DIR / "model.pkl", "rb") as f:
    pipeline = pickle.load(f)

ALLOWED_EXTENSIONS = {".txt", ".pdf", ".docx"}
MIN_TEXT_LENGTH = 20

FRAUD_KEYWORDS = [
    "send id", "bank details", "banking details", "registration fee",
    "no experience needed", "work from home", "guaranteed income",
    "urgent hiring", "send passport", "deposit required", "whatsapp",
    "weekly pay", "daily earnings", "limited slots", "upfront fee",
    "joining fee", "no qualifications needed", "earn r", "admin fee",
    "processing fee", "activation fee", "proof of payment",
]


def extract_text_from_file(file_storage):
    """Extract plain text from an uploaded .txt, .pdf or .docx file.
    Raises ValueError with a user-facing message on any problem."""
    filename = file_storage.filename or ""
    ext = os.path.splitext(filename)[1].lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            "Unsupported file type. Please upload a .txt, .pdf or .docx file."
        )

    data = file_storage.read()
    if not data:
        raise ValueError("The uploaded file is empty.")

    if ext == ".txt":
        try:
            return data.decode("utf-8")
        except UnicodeDecodeError:
            return data.decode("latin-1", errors="ignore")

    if ext == ".pdf":
        import pdfplumber
        try:
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        except Exception:
            raise ValueError("Could not read the PDF file. It may be corrupted.")
        if len(text.strip()) < MIN_TEXT_LENGTH:
            raise ValueError(
                "No readable text found in the PDF. Scanned image PDFs are not "
                "supported - please copy the posting text and paste it instead."
            )
        return text

    if ext == ".docx":
        import docx
        try:
            document = docx.Document(io.BytesIO(data))
        except Exception:
            raise ValueError("Could not read the Word file. It may be corrupted.")
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)


def classify_posting(text):
    proba = pipeline.predict_proba([text])[0]
    fraud_prob = proba[1]
    legit_prob = proba[0]

    if fraud_prob >= 0.60:
        label, risk, css, icon = "Fraudulent", "HIGH RISK", "fraudulent", "danger"
        confidence = round(fraud_prob * 100, 1)
        advice = ("This posting shows strong indicators of fraud. Do not submit "
                  "personal documents, banking details, or pay any fees. Report "
                  "this posting to the job platform immediately.")
    elif fraud_prob >= 0.35:
        label, risk, css, icon = "Suspicious", "MEDIUM RISK", "suspicious", "warning"
        confidence = round(fraud_prob * 100, 1)
        advice = ("This posting contains some characteristics associated with "
                  "fraudulent ads. Proceed with caution and verify the company "
                  "independently before submitting any personal information.")
    else:
        label, risk, css, icon = "Legitimate", "LOW RISK", "legitimate", "safe"
        confidence = round(legit_prob * 100, 1)
        advice = ("This posting appears legitimate based on its content. Always "
                  "verify the company independently and never share banking "
                  "details or pay upfront fees for any job application.")

    lower_text = text.lower()
    found = [kw for kw in FRAUD_KEYWORDS if kw in lower_text]

    return {
        "label": label,
        "risk": risk,
        "css": css,
        "icon": icon,
        "confidence": confidence,
        "fraud_prob": round(fraud_prob * 100, 1),
        "legit_prob": round(legit_prob * 100, 1),
        "indicators": found[:5],
        "advice": advice,
    }


@app.route("/", methods=["GET"])
def index():
    return send_file(BASE_DIR / "index.html")


@app.route("/analyse", methods=["POST"])
def analyse():
    """Single analysis endpoint for both pasted text and uploaded files."""
    source = "pasted text"
    try:
        uploaded = request.files.get("file")
        if uploaded and uploaded.filename:
            text = extract_text_from_file(uploaded)
            source = uploaded.filename
        else:
            data = request.get_json(silent=True) or {}
            text = data.get("text", "")
    except ValueError as err:
        return jsonify({"error": str(err)}), 400

    text = (text or "").strip()
    if len(text) < MIN_TEXT_LENGTH:
        return jsonify({"error": "Please provide a job posting of at least "
                                 f"{MIN_TEXT_LENGTH} characters."}), 400

    result = classify_posting(text)
    result["source"] = source
    result["extracted_text"] = text[:2000]
    return jsonify(result)


@app.route("/api/analyse", methods=["POST"])
def api_analyse():
    """JSON API for programmatic access (browser extensions, third-party tools)."""
    data = request.get_json(silent=True) or {}
    text = data.get("text", "").strip()
    if not text:
        return jsonify({"error": "No text provided"}), 400
    return jsonify(classify_posting(text))


@app.errorhandler(413)
def file_too_large(_):
    return jsonify({"error": "File is too large. The maximum size is 5 MB."}), 413


if __name__ == "__main__":
    app.run(debug=True, port=5000)
